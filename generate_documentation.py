"""
Script to generate comprehensive Word documentation for the Multi-Agent RAG System.
This script creates a detailed Word document explaining every component of the project.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def add_heading_with_style(doc, text, level=1):
    """Add a heading with custom styling."""
    heading = doc.add_heading(text, level=level)
    heading.style.font.size = Pt(16 if level == 1 else 14 if level == 2 else 12)
    return heading

def add_code_block(doc, code, language="python"):
    """Add a code block with monospace font."""
    para = doc.add_paragraph()
    run = para.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 0, 0)
    para.style = 'No Spacing'
    return para

def add_bullet_point(doc, text):
    """Add a bullet point."""
    para = doc.add_paragraph(text, style='List Bullet')
    return para

def create_documentation():
    """Create the comprehensive documentation document."""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title Page
    title = doc.add_heading('Multi-Agent RAG System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Complete Code Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(14)
    
    doc.add_paragraph()
    doc.add_paragraph('This document provides a comprehensive explanation of every component, file, and functionality in the Multi-Agent RAG (Retrieval-Augmented Generation) system.')
    doc.add_paragraph()
    doc.add_paragraph('Generated: December 2025')
    doc.add_page_break()
    
    # Table of Contents
    add_heading_with_style(doc, 'Table of Contents', 1)
    toc_items = [
        '1. Project Overview',
        '2. System Architecture',
        '3. Backend Components',
        '  3.1 Main Application',
        '  3.2 Configuration',
        '  3.3 Agent System',
        '  3.4 Services',
        '  3.5 Repositories',
        '  3.6 API Endpoints',
        '  3.7 Models',
        '  3.8 Utilities',
        '4. Frontend Components',
        '  4.1 Application Structure',
        '  4.2 Pages',
        '  4.3 Components',
        '  4.4 Services',
        '5. Data Flow',
        '6. Key Features',
        '7. Technologies Used',
        '8. Deployment'
    ]
    
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # 1. Project Overview
    add_heading_with_style(doc, '1. Project Overview', 1)
    doc.add_paragraph(
        'The Multi-Agent RAG System is a sophisticated enterprise application that combines Retrieval-Augmented Generation (RAG) with a multi-agent architecture. The system allows users to upload documents (PDFs, DOCX, images with OCR), ask questions about them, and receive intelligent answers from specialized AI agents.'
    )
    
    doc.add_paragraph()
    add_heading_with_style(doc, '1.1 Key Features', 2)
    features = [
        'Multi-Agent Architecture: Specialized agents (IT Policy, HR Policy, Research) coordinated by LangGraph',
        'Intelligent Query Routing: Coordinator agent routes queries to appropriate specialist agents',
        'Document Processing: Supports PDFs, DOCX, TXT, MD, and image files (JPG, PNG, WEBP, BMP, TIFF)',
        'OCR Integration: Extract text from images using EasyOCR',
        'Vector Database: ChromaDB for semantic search and retrieval',
        'Authentication & Authorization: User and admin roles with JWT-based auth',
        'Modern UI: Dark-themed React interface with purple accents',
        'Chat History: Persistent conversation storage with MongoDB',
        'Parallel Execution: LangGraph enables parallel agent processing for improved performance'
    ]
    
    for feature in features:
        add_bullet_point(doc, feature)
    
    doc.add_page_break()
    
    # 2. System Architecture
    add_heading_with_style(doc, '2. System Architecture', 1)
    doc.add_paragraph(
        'The system follows a layered architecture pattern with clear separation of concerns:'
    )
    
    add_heading_with_style(doc, '2.1 Architecture Layers', 2)
    layers = [
        'Presentation Layer: React frontend with modern UI components',
        'API Gateway Layer: FastAPI routers handling HTTP requests',
        'Service Layer: Business logic and orchestration',
        'Data Access Layer: Repositories for MongoDB and ChromaDB',
        'Persistence Layer: MongoDB (Atlas) and ChromaDB (local)'
    ]
    
    for layer in layers:
        add_bullet_point(doc, layer)
    
    add_heading_with_style(doc, '2.2 Agent Workflow', 2)
    doc.add_paragraph(
        'The system uses LangGraph to orchestrate a multi-agent workflow:'
    )
    workflow_steps = [
        'User Query: User submits a question through the frontend',
        'Coordinator Agent: Analyzes the query and determines which specialist agents should handle it',
        'Parallel Agent Execution: Selected agents run simultaneously using asyncio',
        'Document Retrieval: Each agent searches ChromaDB for relevant document chunks',
        'Response Generation: Agents generate answers using OpenAI GPT-4',
        'Synthesis: If multiple agents responded, their answers are synthesized into a final response',
        'Response Delivery: Final answer is returned to the user with source attribution'
    ]
    
    for step in workflow_steps:
        add_bullet_point(doc, step)
    
    doc.add_page_break()
    
    # 3. Backend Components
    add_heading_with_style(doc, '3. Backend Components', 1)
    doc.add_paragraph(
        'The backend is built with FastAPI and Python, organized into clear modules for maintainability.'
    )
    
    # 3.1 Main Application
    add_heading_with_style(doc, '3.1 Main Application (app/main.py)', 2)
    doc.add_paragraph(
        'The main.py file is the entry point of the FastAPI application. It initializes the application, configures middleware, and sets up the API routes.'
    )
    
    doc.add_paragraph('Key Responsibilities:')
    responsibilities = [
        'Creates the FastAPI application instance',
        'Configures CORS middleware for frontend communication',
        'Sets up lifespan events for database connections (startup/shutdown)',
        'Includes all API routers under /api/v1 prefix',
        'Provides root endpoint with API information'
    ]
    
    for resp in responsibilities:
        add_bullet_point(doc, resp)
    
    add_code_block(doc, '''@asynccontextmanager
async def lifespan(app: FastAPI):
    """Lifespan context manager for startup and shutdown events."""
    # Startup: Initialize MongoDB and ChromaDB
    db_repo = await get_mongodb_repo()
    chroma_repo = get_chroma_repo()
    yield
    # Shutdown: Close connections
    await db_repo.disconnect()
    chroma_repo.disconnect()''')
    
    # 3.2 Configuration
    add_heading_with_style(doc, '3.2 Configuration (app/config.py)', 2)
    doc.add_paragraph(
        'The config.py module uses Pydantic Settings to manage all environment variables and application configuration.'
    )
    
    doc.add_paragraph('Configuration Categories:')
    config_categories = [
        'Application Settings: App name, version, debug mode, environment',
        'Server Configuration: Host, port settings',
        'MongoDB Configuration: Connection string, database name, collection names',
        'ChromaDB Configuration: Persist directory, collection name',
        'OpenAI Configuration: API key, model names, temperature, max tokens',
        'JWT Authentication: Secret key, algorithm, token expiration',
        'CORS Configuration: Allowed origins and hosts',
        'Upload Settings: Max file size, allowed extensions',
        'Agent Configuration: Max iterations, timeout settings',
        'OPIK Configuration: Observability and evaluation settings'
    ]
    
    for cat in config_categories:
        add_bullet_point(doc, cat)
    
    doc.add_paragraph()
    doc.add_paragraph(
        'The Settings class uses Pydantic for automatic validation and type safety. All settings can be overridden via environment variables.'
    )
    
    # 3.3 Agent System
    add_heading_with_style(doc, '3.3 Agent System', 2)
    doc.add_paragraph(
        'The agent system is the core of the application, implementing a multi-agent RAG architecture using LangGraph.'
    )
    
    add_heading_with_style(doc, '3.3.1 Base Agent (app/agents/base_agent.py)', 3)
    doc.add_paragraph(
        'BaseAgent is an abstract base class that provides common functionality for all specialist agents.'
    )
    
    doc.add_paragraph('Key Methods:')
    base_methods = [
        'retrieve_relevant_documents(): Queries ChromaDB for relevant document chunks, filters out short chunks (<100 chars)',
        'format_context_from_documents(): Formats retrieved documents into a context string for the LLM',
        'generate_response(): Uses OpenAI GPT-4 to generate a response based on the query and context',
        'process(): Abstract method that must be implemented by subclasses'
    ]
    
    for method in base_methods:
        add_bullet_point(doc, method)
    
    doc.add_paragraph()
    doc.add_paragraph(
        'The BaseAgent includes Opik observability integration for tracing and metrics collection. It filters out junk chunks (headers, footers, page numbers) both at upload time and query time to ensure high-quality search results.'
    )
    
    add_heading_with_style(doc, '3.3.2 Coordinator Agent (app/agents/coordinator_agent.py)', 3)
    doc.add_paragraph(
        'The CoordinatorAgent is responsible for analyzing user queries and routing them to appropriate specialist agents.'
    )
    
    doc.add_paragraph('How It Works:')
    coordinator_steps = [
        'Receives the user query',
        'Uses OpenAI GPT-4 with structured output (Pydantic) to analyze the query',
        'Determines which agents should handle the query (can be multiple)',
        'Returns routing decision with confidence score and reasoning',
        'Validates agent names and provides fallback routing if needed'
    ]
    
    for step in coordinator_steps:
        add_bullet_point(doc, step)
    
    doc.add_paragraph()
    doc.add_paragraph(
        'The coordinator uses a RoutingDecision model with structured output to ensure reliable parsing. It supports routing to multiple agents for complex queries that span multiple domains.'
    )
    
    add_heading_with_style(doc, '3.3.3 Specialist Agents', 3)
    doc.add_paragraph('The system includes three specialist agents:')
    
    doc.add_paragraph('IT Policy Agent (app/agents/it_policy_agent.py):')
    add_bullet_point(doc, 'Handles queries about IT security, infrastructure, software, hardware, network, cybersecurity')
    add_bullet_point(doc, 'Filters documents by document_type="it_policy" in ChromaDB queries')
    add_bullet_point(doc, 'Uses strict system prompt to only answer from available documents')
    
    doc.add_paragraph('HR Policy Agent (app/agents/hr_policy_agent.py):')
    add_bullet_point(doc, 'Handles queries about HR policies, benefits, leave, compensation, onboarding, performance reviews')
    add_bullet_point(doc, 'Filters documents by document_type="hr_policy" in ChromaDB queries')
    add_bullet_point(doc, 'Uses strict system prompt to prevent hallucinations')
    
    doc.add_paragraph('Research Agent (app/agents/research_agent.py):')
    add_bullet_point(doc, 'Handles general research queries that don\'t fit specific domains')
    add_bullet_point(doc, 'Searches across all document types')
    add_bullet_point(doc, 'Used as fallback for ambiguous queries')
    
    add_heading_with_style(doc, '3.3.4 Graph Orchestrator (app/agents/graph_orchestrator.py)', 3)
    doc.add_paragraph(
        'The AgentOrchestrator uses LangGraph to manage the multi-agent workflow with state management and parallel execution.'
    )
    
    doc.add_paragraph('Graph Structure:')
    graph_nodes = [
        'START → coordinator: Entry point',
        'coordinator → execute_agents: Routes to agent execution',
        'execute_agents → synthesize: Executes agents in parallel, then synthesizes',
        'synthesize → END: Final response'
    ]
    
    for node in graph_nodes:
        add_bullet_point(doc, node)
    
    doc.add_paragraph()
    doc.add_paragraph('State Management:')
    doc.add_paragraph(
        'The AgentState TypedDict flows through the graph, containing: query, user_id, session_id, routing_decision, agents_to_invoke, agent_responses (with reducer), final_answer, sources, processing_time, and error.'
    )
    
    doc.add_paragraph()
    doc.add_paragraph('Key Features:')
    orchestrator_features = [
        'Parallel Agent Execution: Uses asyncio.gather() to run multiple agents simultaneously',
        'State Reducers: Uses Annotated types with reducer functions to properly combine agent responses',
        'Opik Observability: Full workflow tracing with confidence scores and metrics',
        'Error Handling: Graceful error handling with fallback responses',
        'Synthesis: LLM-based synthesis when multiple agents respond'
    ]
    
    for feature in orchestrator_features:
        add_bullet_point(doc, feature)
    
    doc.add_page_break()
    
    # 3.4 Services
    add_heading_with_style(doc, '3.4 Services', 2)
    doc.add_paragraph(
        'Services contain the business logic and orchestrate interactions between repositories and agents.'
    )
    
    add_heading_with_style(doc, '3.4.1 Document Service (app/services/document_service.py)', 3)
    doc.add_paragraph(
        'The DocumentService handles all document upload and processing operations.'
    )
    
    doc.add_paragraph('Document Upload Flow:')
    upload_flow = [
        '1. File Validation: Validates file extension and size',
        '2. File Storage: Saves file to disk with unique UUID filename',
        '3. Database Record: Creates document record in MongoDB with status="pending"',
        '4. Text Extraction: Extracts text based on file type (PDF, DOCX, image OCR)',
        '5. Chunking: Splits text into 500-character chunks with 50-character overlap',
        '6. Upload-Time Filtering: Removes chunks < 100 characters (headers, footers)',
        '7. Embedding Generation: Generates vector embeddings using OpenAI text-embedding-3-small',
        '8. ChromaDB Storage: Stores embeddings with metadata in ChromaDB',
        '9. Status Update: Updates MongoDB record with status="completed" and chunk count'
    ]
    
    for step in upload_flow:
        add_bullet_point(doc, step)
    
    doc.add_paragraph()
    doc.add_paragraph('Text Extraction Methods:')
    extraction_methods = [
        'PDF: Uses PyPDF2 to extract text from each page',
        'DOCX: Placeholder for python-docx implementation',
        'Images: Uses EasyOCR to perform Optical Character Recognition',
        'Text Files: Direct file reading for TXT and MD files'
    ]
    
    for method in extraction_methods:
        add_bullet_point(doc, method)
    
    add_heading_with_style(doc, '3.4.2 Chat Service (app/services/chat_service.py)', 3)
    doc.add_paragraph(
        'The ChatService manages chat sessions and message history in MongoDB.'
    )
    
    doc.add_paragraph('Key Operations:')
    chat_operations = [
        'create_session(): Creates a new chat session with title',
        'get_session(): Retrieves a chat session by ID (with authorization)',
        'add_message(): Adds a message (user or assistant) to a session',
        'get_user_sessions(): Lists all sessions for a user with pagination',
        'update_session_title(): Updates the session title',
        'archive_session(): Archives a session',
        'generate_title_from_first_message(): Creates a title from the first user message'
    ]
    
    for op in chat_operations:
        add_bullet_point(doc, op)
    
    add_heading_with_style(doc, '3.4.3 Agent Service (app/services/agent_service.py)', 3)
    doc.add_paragraph(
        'The AgentService provides a high-level interface for the multi-agent system, integrating with evaluation services.'
    )
    
    doc.add_paragraph('Responsibilities:')
    agent_service_resp = [
        'Wraps the AgentOrchestrator for use by API endpoints',
        'Integrates with EvaluationService for response quality assessment',
        'Maps agent names to AgentType enums',
        'Maps confidence levels to ConfidenceLevel enums',
        'Provides agent information endpoint',
        'Handles evaluation metrics and confidence explanations'
    ]
    
    for resp in agent_service_resp:
        add_bullet_point(doc, resp)
    
    add_heading_with_style(doc, '3.4.4 Auth Service (app/services/auth_service.py)', 3)
    doc.add_paragraph(
        'The AuthService handles user authentication and authorization.'
    )
    
    doc.add_paragraph('Features:')
    auth_features = [
        'User Registration: Creates new users with hashed passwords (bcrypt)',
        'User Login: Validates credentials and generates JWT tokens',
        'Token Verification: Validates JWT tokens and extracts user information',
        'Password Hashing: Uses bcrypt for secure password storage',
        'Role-Based Access: Supports user and admin roles'
    ]
    
    for feature in auth_features:
        add_bullet_point(doc, feature)
    
    doc.add_page_break()
    
    # 3.5 Repositories
    add_heading_with_style(doc, '3.5 Repositories', 2)
    doc.add_paragraph(
        'Repositories implement the Repository pattern, providing a clean interface for data access.'
    )
    
    add_heading_with_style(doc, '3.5.1 MongoDB Repository (app/repositories/mongodb_repo.py)', 3)
    doc.add_paragraph(
        'The MongoDBRepository handles all MongoDB operations using Motor (async MongoDB driver).'
    )
    
    doc.add_paragraph('Collections Managed:')
    collections = [
        'Users Collection: User accounts, authentication data, roles',
        'Chat History Collection: Chat sessions and messages',
        'Documents Collection: Document metadata, status, processing information'
    ]
    
    for coll in collections:
        add_bullet_point(doc, coll)
    
    doc.add_paragraph()
    doc.add_paragraph('Key Operations:')
    mongo_ops = [
        'User Operations: create_user(), get_user_by_email(), get_user_by_id(), update_last_login()',
        'Chat Operations: create_chat_session(), get_chat_session(), add_message_to_session(), get_user_chat_sessions(), update_session_title(), archive_session()',
        'Document Operations: create_document(), get_document(), update_document_status(), get_user_documents(), delete_document()',
        'Index Creation: Automatically creates indexes on startup for optimal query performance'
    ]
    
    for op in mongo_ops:
        add_bullet_point(doc, op)
    
    add_heading_with_style(doc, '3.5.2 ChromaDB Repository (app/repositories/chroma_repo.py)', 3)
    doc.add_paragraph(
        'The ChromaRepository manages vector database operations for document embeddings.'
    )
    
    doc.add_paragraph('Key Operations:')
    chroma_ops = [
        'connect(): Establishes connection to persistent ChromaDB instance',
        'add_documents(): Stores document chunks with embeddings and metadata (batched for large documents)',
        'query_documents(): Performs vector similarity search with metadata filtering',
        'query_by_text(): Convenience method that formats query results',
        'delete_documents(): Deletes documents by IDs or metadata filters',
        'get_collection_count(): Returns the number of documents in the collection',
        'peek(): Samples documents from the collection'
    ]
    
    for op in chroma_ops:
        add_bullet_point(doc, op)
    
    doc.add_paragraph()
    doc.add_paragraph(
        'The repository handles batching for large document uploads (max 1000 documents per batch) to avoid ChromaDB limits. It supports metadata filtering for domain-specific searches (e.g., document_type="hr_policy").'
    )
    
    doc.add_page_break()
    
    # 3.6 API Endpoints
    add_heading_with_style(doc, '3.6 API Endpoints', 2)
    doc.add_paragraph(
        'API endpoints are organized into separate router modules for maintainability.'
    )
    
    add_heading_with_style(doc, '3.6.1 Authentication Endpoints (app/api/auth.py)', 3)
    doc.add_paragraph('Endpoints:')
    auth_endpoints = [
        'POST /api/v1/auth/register: Register a new user account',
        'POST /api/v1/auth/login: User login (returns JWT token)',
        'POST /api/v1/auth/admin/login: Admin login',
        'GET /api/v1/auth/me: Get current user information',
        'GET /api/v1/auth/verify-admin: Verify admin access'
    ]
    
    for endpoint in auth_endpoints:
        add_bullet_point(doc, endpoint)
    
    add_heading_with_style(doc, '3.6.2 Chat Endpoints (app/api/chat.py)', 3)
    doc.add_paragraph('Endpoints:')
    chat_endpoints = [
        'POST /api/v1/chat/query: Process a chat query through the multi-agent system',
        'GET /api/v1/chat/sessions: List all chat sessions for the current user',
        'GET /api/v1/chat/sessions/{session_id}: Get a specific chat session',
        'PUT /api/v1/chat/sessions/{session_id}/title: Update session title',
        'DELETE /api/v1/chat/sessions/{session_id}: Archive a session',
        'GET /api/v1/chat/agents/info: Get information about available agents'
    ]
    
    for endpoint in chat_endpoints:
        add_bullet_point(doc, endpoint)
    
    doc.add_paragraph()
    doc.add_paragraph(
        'The /chat/query endpoint is the main entry point for user queries. It creates or retrieves a chat session, adds the user message, processes it through the agent system, and saves the assistant response.'
    )
    
    add_heading_with_style(doc, '3.6.3 Document Endpoints (app/api/documents.py)', 3)
    doc.add_paragraph('Endpoints:')
    doc_endpoints = [
        'POST /api/v1/documents/upload: Upload and process a document',
        'GET /api/v1/documents/: List user\'s documents with pagination',
        'GET /api/v1/documents/{id}: Get document details',
        'DELETE /api/v1/documents/{id}: Delete a document (admin only)'
    ]
    
    for endpoint in doc_endpoints:
        add_bullet_point(doc, endpoint)
    
    add_heading_with_style(doc, '3.6.4 Health Endpoints (app/api/health.py)', 3)
    doc.add_paragraph('Endpoints:')
    health_endpoints = [
        'GET /api/v1/health: Basic health check',
        'GET /api/v1/health/detailed: Detailed health check with database status'
    ]
    
    for endpoint in health_endpoints:
        add_bullet_point(doc, endpoint)
    
    # 3.7 Models
    add_heading_with_style(doc, '3.7 Models (app/models/)', 2)
    doc.add_paragraph(
        'Models use Pydantic for data validation and serialization.'
    )
    
    doc.add_paragraph('Model Files:')
    model_files = [
        'user.py: User, UserCreate models for authentication',
        'chat.py: ChatSession, ChatMessage, ChatRequest, ChatResponse, AgentType, ConfidenceLevel, EvaluationMetrics',
        'document.py: Document, DocumentUpload, DocumentResponse, DocumentStatus, DocumentType'
    ]
    
    for model_file in model_files:
        add_bullet_point(doc, model_file)
    
    # 3.8 Utilities
    add_heading_with_style(doc, '3.8 Utilities (app/utils/)', 2)
    
    add_heading_with_style(doc, '3.8.1 Embeddings (app/utils/embeddings.py)', 3)
    doc.add_paragraph(
        'The EmbeddingGenerator handles text chunking and embedding generation using OpenAI\'s text-embedding-3-small model.'
    )
    
    doc.add_paragraph('Key Methods:')
    embedding_methods = [
        'chunk_text(): Splits text into overlapping chunks (500 chars, 50 overlap)',
        'generate_embeddings_openai(): Generates embeddings using OpenAI API (batch processing)',
        'get_embedding_generator(): Factory function for dependency injection'
    ]
    
    for method in embedding_methods:
        add_bullet_point(doc, method)
    
    add_heading_with_style(doc, '3.8.2 JWT Handler (app/utils/jwt_handler.py)', 3)
    doc.add_paragraph(
        'Handles JWT token creation, validation, and user extraction.'
    )
    
    add_heading_with_style(doc, '3.8.3 Observability (app/utils/observability.py)', 3)
    doc.add_paragraph(
        'Integrates with Opik for LLM observability, tracing, and evaluation metrics.'
    )
    
    add_heading_with_style(doc, '3.8.4 Validators (app/utils/validators.py)', 3)
    doc.add_paragraph(
        'Contains validation functions for file uploads, extensions, and sizes.'
    )
    
    doc.add_page_break()
    
    # 4. Frontend Components
    add_heading_with_style(doc, '4. Frontend Components', 1)
    doc.add_paragraph(
        'The frontend is built with React and Vite, providing a modern, responsive user interface.'
    )
    
    add_heading_with_style(doc, '4.1 Application Structure', 2)
    doc.add_paragraph('Frontend Structure:')
    frontend_structure = [
        'src/App.jsx: Main application component with routing',
        'src/main.jsx: Application entry point',
        'src/pages/: Page components (UserDashboard, AdminDashboard, Login, Register)',
        'src/components/: Reusable UI components (Chat, DocumentUpload, Navbar, ProtectedRoute)',
        'src/services/api.js: API client with axios interceptors',
        'src/context/AuthContext.jsx: Authentication context provider',
        'src/styles/global.css: Global styles with dark theme'
    ]
    
    for item in frontend_structure:
        add_bullet_point(doc, item)
    
    add_heading_with_style(doc, '4.2 Pages', 2)
    
    doc.add_paragraph('UserDashboard (src/pages/UserDashboard.jsx):')
    doc.add_paragraph(
        'Main dashboard for regular users. Contains the chat interface and chat history sidebar. Manages session state and user profile display.'
    )
    
    doc.add_paragraph('AdminDashboard (src/pages/AdminDashboard.jsx):')
    doc.add_paragraph(
        'Admin-only dashboard for document management and system administration.'
    )
    
    doc.add_paragraph('UserLogin (src/pages/UserLogin.jsx):')
    doc.add_paragraph(
        'Login page for regular users with email/password authentication.'
    )
    
    doc.add_paragraph('AdminLogin (src/pages/AdminLogin.jsx):')
    doc.add_paragraph(
        'Login page for administrators.'
    )
    
    doc.add_paragraph('UserRegister (src/pages/UserRegister.jsx):')
    doc.add_paragraph(
        'User registration page with form validation.'
    )
    
    add_heading_with_style(doc, '4.3 Components', 2)
    
    doc.add_paragraph('ChatInterface (src/components/Chat/ChatInterface.jsx):')
    doc.add_paragraph(
        'Main chat interface component. Handles message sending, receiving, and display. Manages session creation and message history loading.'
    )
    
    doc.add_paragraph('ChatHistory (src/components/Chat/ChatHistory.jsx):')
    doc.add_paragraph(
        'Sidebar component displaying list of chat sessions. Allows session selection and creation of new sessions.'
    )
    
    doc.add_paragraph('MessageBubble (src/components/Chat/MessageBubble.jsx):')
    doc.add_paragraph(
        'Individual message display component. Shows user and assistant messages with different styling. Displays source attribution for assistant messages.'
    )
    
    doc.add_paragraph('DocumentUpload (src/components/DocumentUpload.jsx):')
    doc.add_paragraph(
        'Document upload component with drag-and-drop support. Handles file selection, document type selection, and upload progress.'
    )
    
    doc.add_paragraph('ProtectedRoute (src/components/ProtectedRoute.jsx):')
    doc.add_paragraph(
        'Route guard component that protects routes requiring authentication. Redirects to login if user is not authenticated.'
    )
    
    add_heading_with_style(doc, '4.4 Services', 2)
    
    doc.add_paragraph('API Service (src/services/api.js):')
    doc.add_paragraph(
        'Centralized API client using axios. Includes request/response interceptors for JWT token management and automatic token refresh.'
    )
    
    doc.add_paragraph('API Modules:')
    api_modules = [
        'authAPI: Authentication endpoints (login, register, getCurrentUser)',
        'chatAPI: Chat endpoints (sendMessage, getSessions, getSession, updateSessionTitle, archiveSession)',
        'documentAPI: Document endpoints (upload, list, get, delete)',
        'healthAPI: Health check endpoints'
    ]
    
    for module in api_modules:
        add_bullet_point(doc, module)
    
    doc.add_paragraph()
    doc.add_paragraph(
        'The API service automatically adds JWT tokens to requests and handles token refresh on 401 errors.'
    )
    
    doc.add_page_break()
    
    # 5. Data Flow
    add_heading_with_style(doc, '5. Data Flow', 1)
    
    add_heading_with_style(doc, '5.1 Document Upload Flow', 2)
    doc.add_paragraph(
        'When a user uploads a document, the following flow occurs:'
    )
    
    upload_flow_detailed = [
        '1. Frontend sends file to POST /api/v1/documents/upload',
        '2. DocumentService validates file (extension, size)',
        '3. File is saved to disk with UUID filename',
        '4. MongoDB document record created with status="pending"',
        '5. Text extraction based on file type (PDF, DOCX, OCR)',
        '6. Text is chunked into 500-character pieces with 50-character overlap',
        '7. Short chunks (<100 chars) are filtered out',
        '8. OpenAI embeddings generated for remaining chunks',
        '9. Embeddings stored in ChromaDB with metadata',
        '10. MongoDB record updated with status="completed" and chunk count'
    ]
    
    for step in upload_flow_detailed:
        add_bullet_point(doc, step)
    
    add_heading_with_style(doc, '5.2 Query Processing Flow', 2)
    doc.add_paragraph(
        'When a user submits a query, the following flow occurs:'
    )
    
    query_flow_detailed = [
        '1. Frontend sends query to POST /api/v1/chat/query',
        '2. ChatService creates or retrieves chat session',
        '3. User message added to session in MongoDB',
        '4. AgentService processes query through orchestrator',
        '5. CoordinatorAgent analyzes query and routes to agents',
        '6. Selected agents execute in parallel:',
        '   a. Generate query embedding using OpenAI',
        '   b. Query ChromaDB for relevant chunks (with domain filtering)',
        '   c. Filter out short chunks (<100 chars)',
        '   d. Generate response using GPT-4 with context',
        '7. If multiple agents responded, synthesis node combines answers',
        '8. EvaluationService evaluates response quality',
        '9. Assistant message added to session with metadata',
        '10. Response returned to frontend with sources and confidence'
    ]
    
    for step in query_flow_detailed:
        add_bullet_point(doc, step)
    
    doc.add_page_break()
    
    # 6. Key Features
    add_heading_with_style(doc, '6. Key Features Explained', 1)
    
    add_heading_with_style(doc, '6.1 Multi-Agent Architecture', 2)
    doc.add_paragraph(
        'The system uses a coordinator-specialist pattern where a coordinator agent routes queries to specialized agents. This allows for domain-specific expertise while maintaining flexibility for general queries.'
    )
    
    add_heading_with_style(doc, '6.2 Two-Layer Filtering Strategy', 2)
    doc.add_paragraph(
        'To ensure high-quality search results, the system filters out junk chunks (headers, footers, page numbers) at two points:'
    )
    
    filtering_layers = [
        'Upload-Time Filtering: Removes chunks < 100 characters before embedding. This saves API costs and storage space.',
        'Query-Time Filtering: Retrieves 100 results, filters out short chunks, then returns top 5. This handles legacy documents uploaded before filtering was implemented.'
    ]
    
    for layer in filtering_layers:
        add_bullet_point(doc, layer)
    
    add_heading_with_style(doc, '6.3 Parallel Agent Execution', 2)
    doc.add_paragraph(
        'When multiple agents are selected by the coordinator, they execute in parallel using Python\'s asyncio.gather(). This significantly reduces response time compared to sequential execution.'
    )
    
    add_heading_with_style(doc, '6.4 Opik Observability', 2)
    doc.add_paragraph(
        'The system integrates with Opik for LLM observability, providing:'
    )
    
    observability_features = [
        'Workflow Tracing: Complete trace of query processing through the agent system',
        'Confidence Tracking: Routing confidence scores and quality metrics',
        'Performance Metrics: Latency tracking for each agent and the overall workflow',
        'Evaluation Metrics: Hallucination detection, answer relevance, context precision/recall'
    ]
    
    for feature in observability_features:
        add_bullet_point(doc, feature)
    
    add_heading_with_style(doc, '6.5 Strict Answer Generation', 2)
    doc.add_paragraph(
        'All agents use strict system prompts that instruct the LLM to only answer from the provided context. If information is not found in the documents, agents explicitly state that the information is not available.'
    )
    
    doc.add_page_break()
    
    # 7. Technologies Used
    add_heading_with_style(doc, '7. Technologies Used', 1)
    
    add_heading_with_style(doc, '7.1 Backend Technologies', 2)
    backend_tech = [
        'FastAPI: Modern, high-performance web framework for building APIs',
        'LangGraph: Stateful multi-agent orchestration framework',
        'LangChain: Individual agent implementations and LLM integration',
        'OpenAI API: GPT-4 for LLM responses and text-embedding-3-small for embeddings',
        'ChromaDB: Vector database for semantic search (local persistent storage)',
        'MongoDB: NoSQL database for chat history and user data (MongoDB Atlas)',
        'Motor: Async MongoDB driver for Python',
        'EasyOCR: Optical Character Recognition for image text extraction',
        'PyPDF2: PDF text extraction',
        'Pydantic: Data validation and settings management',
        'JWT: JSON Web Tokens for authentication',
        'bcrypt: Password hashing',
        'Opik: LLM observability and evaluation'
    ]
    
    for tech in backend_tech:
        add_bullet_point(doc, tech)
    
    add_heading_with_style(doc, '7.2 Frontend Technologies', 2)
    frontend_tech = [
        'React.js: UI library for building interactive interfaces',
        'React Router: Client-side routing',
        'Vite: Fast build tool and dev server',
        'Axios: HTTP client for API requests',
        'Lucide React: Icon library',
        'CSS3: Modern styling with glassmorphism effects'
    ]
    
    for tech in frontend_tech:
        add_bullet_point(doc, tech)
    
    doc.add_page_break()
    
    # 8. Deployment
    add_heading_with_style(doc, '8. Deployment and Configuration', 1)
    
    add_heading_with_style(doc, '8.1 Environment Variables', 2)
    doc.add_paragraph('Required Environment Variables:')
    
    env_vars = [
        'OPENAI_API_KEY: OpenAI API key for LLM and embeddings',
        'MONGODB_URL: MongoDB Atlas connection string',
        'JWT_SECRET_KEY: Secret key for JWT token signing',
        'OPIK_API_KEY: (Optional) Opik API key for observability',
        'OPIK_WORKSPACE: (Optional) Opik workspace name',
        'OPIK_PROJECT_NAME: (Optional) Opik project name'
    ]
    
    for var in env_vars:
        add_bullet_point(doc, var)
    
    add_heading_with_style(doc, '8.2 Running the Application', 2)
    
    doc.add_paragraph('Backend:')
    add_code_block(doc, 'cd backend\npython -m uvicorn app.main:app --reload --host 0.0.0.0 --port 8000')
    
    doc.add_paragraph('Frontend:')
    add_code_block(doc, 'cd frontend\nnpm install\nnpm start')
    
    add_heading_with_style(doc, '8.3 Database Setup', 2)
    doc.add_paragraph(
        'MongoDB: Requires a MongoDB Atlas account (or local MongoDB). The connection string should be provided in MONGODB_URL.'
    )
    doc.add_paragraph(
        'ChromaDB: Automatically creates persistent storage in ./data/chroma on first run. No additional setup required.'
    )
    
    # Conclusion
    doc.add_page_break()
    add_heading_with_style(doc, 'Conclusion', 1)
    doc.add_paragraph(
        'This Multi-Agent RAG System represents a sophisticated implementation of retrieval-augmented generation with multi-agent orchestration. The system is designed for scalability, maintainability, and high-quality responses through strict answer generation, intelligent routing, and comprehensive observability.'
    )
    
    doc.add_paragraph()
    doc.add_paragraph(
        'Key architectural decisions include:'
    )
    
    architectural_decisions = [
        'Separation of concerns with clear service and repository layers',
        'Multi-agent architecture for domain-specific expertise',
        'Two-layer filtering for high-quality search results',
        'Parallel agent execution for performance',
        'Comprehensive observability for monitoring and debugging',
        'Strict answer generation to prevent hallucinations'
    ]
    
    for decision in architectural_decisions:
        add_bullet_point(doc, decision)
    
    doc.add_paragraph()
    doc.add_paragraph(
        'The system is production-ready with proper error handling, authentication, authorization, and data persistence.'
    )
    
    # Save document
    output_path = 'Multi-Agent_RAG_System_Documentation.docx'
    doc.save(output_path)
    print(f"Documentation generated successfully: {output_path}")
    return output_path

if __name__ == "__main__":
    create_documentation()

